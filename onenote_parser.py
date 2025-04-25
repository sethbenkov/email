import os
import glob
from datetime import datetime
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from dotenv import load_dotenv

import config

def find_latest_docx(folder_path):
    """Finds the most recently modified .docx file in a folder."""
    if not os.path.isdir(folder_path):
        print(f"Error: OneNote export folder not found: {folder_path}")
        return None
    
    search_pattern = os.path.join(folder_path, '*.docx')
    docx_files = glob.glob(search_pattern)
    
    if not docx_files:
        print(f"No .docx files found in {folder_path}")
        return None
    
    try:
        latest_file = max(docx_files, key=os.path.getmtime)
        return latest_file
    except Exception as e:
        print(f"Error finding latest file in {folder_path}: {e}")
        return None

def parse_onenote_docx(file_path):
    """Parses a .docx file exported from OneNote, extracting tasks not marked with DONE."""
    tasks = []
    try:
        document = Document(file_path)
        print(f"Parsing OneNote export: {os.path.basename(file_path)}")
        
        # Iterate through paragraphs (simplest approach)
        # Might need refinement based on how OneNote structures lists/tasks in DOCX
        for para in document.paragraphs:
            text = para.text.strip()
            if text: # Ignore empty paragraphs
                # Check if the line starts with the DONE marker (case-insensitive)
                if not text.upper().startswith(config.ONENOTE_DONE_MARKER):
                    # Check if it looks like a list item (starts with bullet, number, etc.)
                    # This helps filter out regular paragraphs if needed, but might be too strict.
                    # For now, let's include any non-DONE line.
                    # if re.match(r'^\s*[-*\u2022\u25E6\u25CF]|\d+\.\s', text):
                    print(f"  * [Found Task] {text}")
                    tasks.append(text)
                # else:
                    # print(f"  - [Skipped DONE] {text}")
        
        # Alternative: Iterate through tables if tasks are in tables
        # for table in document.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             # Process cell text similarly
        #             pass
            
    except PackageNotFoundError:
        print(f"Error: Could not open file {file_path}. It might be corrupted or not a valid .docx file.")
        return ["Error: Invalid OneNote export file found."]
    except Exception as e:
        print(f"Error parsing Word document {file_path}: {e}")
        return [f"Error parsing OneNote export: {e}"]

    if not tasks:
        print("No open tasks found in the document.")
        return ["No open tasks found in latest OneNote export."]
        
    return tasks

def get_onenote_tasks_from_export():
    """Finds the latest OneNote .docx export and parses it for open tasks."""
    print("\n--- Parsing OneNote Export File --- ")
    load_dotenv()
    export_folder = os.getenv(config.ONENOTE_EXPORT_FOLDER_ENV_VAR)

    if not export_folder:
        print(f"Error: Environment variable {config.ONENOTE_EXPORT_FOLDER_ENV_VAR} not set in .env file.")
        return [f"Error: OneNote export folder path not configured."]
    
    latest_export_file = find_latest_docx(export_folder)
    
    if not latest_export_file:
        return ["Error: Could not find a recent OneNote .docx export file."]
        
    tasks = parse_onenote_docx(latest_export_file)
    return tasks

if __name__ == '__main__':
    # For local testing
    # Ensure you have a .docx file in the folder specified in your .env
    tasks = get_onenote_tasks_from_export()
    print("\nOpen OneNote Tasks from DOCX:")
    for task in tasks:
        print(task) 